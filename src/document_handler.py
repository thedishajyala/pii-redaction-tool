from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph


class DocumentHandler:
    """
    Handles loading and saving Word documents.
    """

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document: DocumentType = Document(file_path)

    def get_paragraphs(self) -> list[Paragraph]:
        """Return all paragraphs."""
        return self.document.paragraphs

    def get_tables(self) -> list[Table]:
        """Return all tables."""
        return self.document.tables

    def save(self, output_path: str):
        """Save document."""
        self.document.save(output_path)
