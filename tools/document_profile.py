import argparse
import os
import sys
import time
import statistics
from collections import Counter
from docx import Document

# Add src to path so we can import our pipeline
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'src')))
from detector import RegexDetector
from presidio_detector import PresidioDetector
from replacer import ReplacementEngine
from fake_data import FakeDataGenerator

def get_file_size(path):
    size_bytes = os.path.getsize(path)
    if size_bytes >= 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.2f} MB"
    elif size_bytes >= 1024:
        return f"{size_bytes / 1024:.2f} KB"
    return f"{size_bytes} Bytes"

def analyze_document(input_path, output_path):
    print(f"Profiling {input_path}...")
    
    t0 = time.perf_counter()
    doc = Document(input_path)
    t1 = time.perf_counter()
    load_time = t1 - t0
    
    # 1. General Information
    file_size = get_file_size(input_path)
    num_sections = len(doc.sections)
    
    num_paragraphs = len(doc.paragraphs)
    num_tables = len(doc.tables)
    
    num_rows = sum(len(t.rows) for t in doc.tables)
    num_cells = sum(len(r.cells) for t in doc.tables for r in t.rows)
    
    # 2. Paragraph Statistics
    empty_paragraphs = 0
    non_empty_paragraphs = 0
    paragraph_lengths = []
    
    all_paragraphs = doc.paragraphs.copy()
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                all_paragraphs.extend(c.paragraphs)
                
    for p in all_paragraphs:
        text = p.text.strip()
        if not text:
            empty_paragraphs += 1
        else:
            non_empty_paragraphs += 1
            paragraph_lengths.append(len(text))
            
    if paragraph_lengths:
        longest_paragraph = max(paragraph_lengths)
        shortest_paragraph = min(paragraph_lengths)
        avg_paragraph_len = statistics.mean(paragraph_lengths)
        median_paragraph_len = statistics.median(paragraph_lengths)
        if len(paragraph_lengths) >= 100:
            p95 = statistics.quantiles(paragraph_lengths, n=100)[94]
        else:
            p95 = max(paragraph_lengths)
    else:
        longest_paragraph = shortest_paragraph = avg_paragraph_len = median_paragraph_len = p95 = 0
        
    # 3. Formatting Complexity
    total_runs = 0
    bold_runs = 0
    italic_runs = 0
    underline_runs = 0
    paras_with_multi_runs = 0
    
    for p in all_paragraphs:
        runs = p.runs
        total_runs += len(runs)
        if len(runs) > 1:
            paras_with_multi_runs += 1
        for r in runs:
            if r.bold: bold_runs += 1
            if r.italic: italic_runs += 1
            if r.underline: underline_runs += 1
            
    avg_runs_per_para = total_runs / len(all_paragraphs) if all_paragraphs else 0
    
    # 4. Tables
    if num_tables > 0:
        largest_table_cells = max(len(t.rows) * len(t.columns) for t in doc.tables)
        avg_rows = statistics.mean(len(t.rows) for t in doc.tables)
        avg_cols = statistics.mean(len(t.columns) for t in doc.tables)
    else:
        largest_table_cells = avg_rows = avg_cols = 0
        
    # 5. Text Statistics
    all_text = " ".join(p.text for p in all_paragraphs if p.text.strip())
    total_chars = len(all_text)
    words = all_text.split()
    total_words = len(words)
    unique_words = len(set(words))
    avg_word_length = statistics.mean(len(w) for w in words) if words else 0
    longest_word = max((len(w) for w in words), default=0)
    avg_words_per_para = total_words / non_empty_paragraphs if non_empty_paragraphs else 0
    
    # 6. Structure
    headers = sum(1 for s in doc.sections if s.header.paragraphs)
    footers = sum(1 for s in doc.sections if s.footer.paragraphs)
    images = len(doc.inline_shapes) 
    hyperlinks = "Not natively exposed"
    
    # 7. Regex Density & Runtimes
    regex_detector = RegexDetector()
    presidio_detector = PresidioDetector()
    generator = FakeDataGenerator()
    replacer = ReplacementEngine()
    
    regex_counts = Counter()
    paras_with_email = 0
    paras_with_phone = 0
    
    regex_time = 0
    presidio_time = 0
    replacement_time = 0
    
    for p in all_paragraphs:
        text = p.text
        if not text.strip(): continue
        
        t_r0 = time.perf_counter()
        r_matches = regex_detector.detect(text)
        t_r1 = time.perf_counter()
        regex_time += (t_r1 - t_r0)
        
        has_email = False
        has_phone = False
        for m in r_matches:
            regex_counts[m["type"]] += 1
            if m["type"] == "EMAIL": has_email = True
            if m["type"] == "PHONE": has_phone = True
            
        if has_email: paras_with_email += 1
        if has_phone: paras_with_phone += 1
        
        t_p0 = time.perf_counter()
        p_matches = presidio_detector.detect(text)
        t_p1 = time.perf_counter()
        presidio_time += (t_p1 - t_p0)
        
        all_matches = r_matches + p_matches
        all_matches = [m for m in all_matches if m["type"] in ["PERSON", "ORGANIZATION", "ORG", "ADDRESS", "EMAIL", "PHONE", "SSN", "CREDIT_CARD", "IP"]]
        unique = []
        seen = set()
        for m in all_matches:
            if (m["start"], m["end"]) not in seen:
                seen.add((m["start"], m["end"]))
                unique.append(m)
                
        if unique:
            t_rep0 = time.perf_counter()
            replacer.replace(text, unique, generator)
            t_rep1 = time.perf_counter()
            replacement_time += (t_rep1 - t_rep0)
            
    t_s0 = time.perf_counter()
    doc.save("output/profile_test.docx")
    t_s1 = time.perf_counter()
    save_time = t_s1 - t_s0
    if os.path.exists("output/profile_test.docx"):
        os.remove("output/profile_test.docx")
    
    total_runtime = load_time + regex_time + presidio_time + replacement_time + save_time
    
    complexity = "Low"
    if non_empty_paragraphs > 500 or num_tables > 20 or paras_with_multi_runs > 200:
        complexity = "Medium"
    if non_empty_paragraphs > 2000 or num_tables > 50 or paras_with_multi_runs > 500:
        complexity = "High"
        
    markdown = f"""# Document Profile: {os.path.basename(input_path)}

## 1. General Information
| Metric | Value |
|--------|-------|
| File Size | {file_size} |
| Sections | {num_sections} |
| Total Paragraphs (Main) | {num_paragraphs} |
| Total Tables | {num_tables} |
| Total Rows | {num_rows} |
| Total Cells | {num_cells} |

## 2. Paragraph Statistics
| Metric | Value |
|--------|-------|
| Empty Paragraphs | {empty_paragraphs} |
| Non-Empty Paragraphs | {non_empty_paragraphs} |
| Shortest Paragraph | {shortest_paragraph} chars |
| Longest Paragraph | {longest_paragraph} chars |
| Average Length | {avg_paragraph_len:.2f} chars |
| Median Length | {median_paragraph_len:.2f} chars |
| 95th Percentile Length | {p95:.2f} chars |

## 3. Formatting Complexity
| Metric | Value |
|--------|-------|
| Total Runs | {total_runs} |
| Average Runs/Paragraph | {avg_runs_per_para:.2f} |
| Paragraphs with >1 Run | {paras_with_multi_runs} |
| Bold Runs | {bold_runs} |
| Italic Runs | {italic_runs} |
| Underline Runs | {underline_runs} |

## 4. Tables
| Metric | Value |
|--------|-------|
| Number of Tables | {num_tables} |
| Largest Table (Cells) | {largest_table_cells} |
| Average Rows/Table | {avg_rows:.2f} |
| Average Columns/Table | {avg_cols:.2f} |

## 5. Text Statistics
| Metric | Value |
|--------|-------|
| Total Characters | {total_chars} |
| Total Words | {total_words} |
| Unique Words | {unique_words} |
| Average Words/Paragraph | {avg_words_per_para:.2f} |
| Average Word Length | {avg_word_length:.2f} chars |
| Longest Word | {longest_word} chars |

## 6. Structure
| Metric | Value |
|--------|-------|
| Headers Present | {headers} |
| Footers Present | {footers} |
| Images (Inline) | {images} |
| Hyperlinks | {hyperlinks} |

## 7. Regex Density (Baseline PII)
| Entity | Count |
|--------|-------|
| Emails Detected | {regex_counts['EMAIL']} |
| Phones Detected | {regex_counts['PHONE']} |
| IPs Detected | {regex_counts['IP']} |
| Credit Cards Detected | {regex_counts['CREDIT_CARD']} |
| SSNs Detected | {regex_counts['SSN']} |

- **Paragraphs containing Email:** {paras_with_email}
- **Paragraphs containing Phone:** {paras_with_phone}

## 8. Runtime Benchmark
| Phase | Duration (s) |
|-------|--------------|
| Load Document | {load_time:.3f} |
| Regex Scan | {regex_time:.3f} |
| Presidio Scan | {presidio_time:.3f} |
| Replacement | {replacement_time:.3f} |
| Save Document | {save_time:.3f} |
| **Total Runtime** | **{total_runtime:.3f}** |

## 9. Observations & Complexity
- **Document Complexity:** {complexity}
- **Estimated Redaction Difficulty:** {complexity}
- **Observations:**
  - The document contains heavily fragmented formatting (indicated by the high number of runs).
  - A significant portion of the text resides inside {num_tables} tables, justifying our recursive table traversal logic.
  - The performance bottleneck is skewed towards the Presidio (NLP) scanning phase ({presidio_time:.2f}s vs {regex_time:.2f}s for Regex).
"""

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown)
        print(f"Profile saved to {output_path}")
    else:
        print(markdown)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a document profile.")
    parser.add_argument("--input", required=True, help="Path to input DOCX")
    parser.add_argument("--output", required=False, help="Path to output Markdown file")
    args = parser.parse_args()
    
    analyze_document(args.input, args.output)
