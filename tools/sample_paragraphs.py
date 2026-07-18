import sys
import os
import random
import argparse
from collections import defaultdict
from docx import Document

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'src')))
from detector import RegexDetector
from presidio_detector import PresidioDetector

# From our fix in main.py
SUPPORTED_ENTITY_TYPES = {
    "PERSON",
    "ORGANIZATION",
    "ORG",
    "ADDRESS",
    "EMAIL",
    "PHONE",
    "SSN",
    "CREDIT_CARD",
    "IP"
}

def generate_sample(input_path, output_path):
    print(f"Loading document {input_path}...")
    doc = Document(input_path)
    
    regex_detector = RegexDetector()
    presidio_detector = PresidioDetector()
    
    # Extract all paragraphs and their metadata
    all_paras = []
    
    pid = 0
    for p in doc.paragraphs:
        if not p.text.strip(): continue
        pid += 1
        all_paras.append({
            "pid": pid,
            "text": p.text.strip(),
            "runs_count": len(p.runs),
            "is_table": False
        })
        
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if not p.text.strip(): continue
                    pid += 1
                    all_paras.append({
                        "pid": pid,
                        "text": p.text.strip(),
                        "runs_count": len(p.runs),
                        "is_table": True
                    })
                    
    print(f"Total non-empty paragraphs to scan: {len(all_paras)}")
    
    # Bins
    categories = defaultdict(list)
    
    for i, item in enumerate(all_paras):
        if i % 1000 == 0:
            print(f"Scanned {i} paragraphs...")
            
        text = item["text"]
        r_matches = regex_detector.detect(text)
        p_matches = presidio_detector.detect(text)
        all_matches = r_matches + p_matches
        all_matches = [m for m in all_matches if m["type"] in SUPPORTED_ENTITY_TYPES]
        
        # Dedupe
        unique = []
        seen = set()
        for m in all_matches:
            identifier = (m["start"], m["end"])
            if identifier not in seen:
                seen.add(identifier)
                unique.append(m)
                
        item["matches"] = unique
        
        has_pii = False
        types = set(m["type"] for m in unique)
        
        if "EMAIL" in types:
            categories["emails"].append(item)
            has_pii = True
        if "PHONE" in types:
            categories["phones"].append(item)
            has_pii = True
        if "PERSON" in types:
            categories["persons"].append(item)
            has_pii = True
        if "ORGANIZATION" in types or "ORG" in types:
            categories["organizations"].append(item)
            has_pii = True
        if "ADDRESS" in types:
            categories["addresses"].append(item)
            has_pii = True
            
        if not has_pii:
            categories["no_pii"].append(item)
            
        if item["is_table"]:
            categories["tables"].append(item)
            
    # Sample based on quotas
    quotas = {
        "emails": 20,
        "phones": 15,
        "persons": 20,
        "organizations": 15,
        "addresses": 10,
        "tables": 10,
        "no_pii": 10
    }
    
    selected = {} # Use dict to avoid duplicates by pid
    
    for cat, quota in quotas.items():
        candidates = categories[cat]
        available = [c for c in candidates if c["pid"] not in selected]
        
        k = min(quota, len(available))
        if k > 0:
            sampled = random.sample(available, k)
            for item in sampled:
                selected[item["pid"]] = item
                
    # Top up to 100 if needed
    shortfall = 100 - len(selected)
    if shortfall > 0:
        remaining = [c for c in all_paras if c["pid"] not in selected]
        if remaining:
            top_up = random.sample(remaining, min(shortfall, len(remaining)))
            for item in top_up:
                selected[item["pid"]] = item
                
    final_sample = sorted(selected.values(), key=lambda x: x["pid"])
    print(f"\nFinal sampled set size: {len(final_sample)}")
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("# Ground Truth Annotations (Draft)\n\n")
        f.write("> **Instructions**: Review each paragraph. For the 'Entities:' section, accept/reject or modify the list. Format is `- [TYPE] Value`.\n\n")
        
        for item in final_sample:
            f.write(f"### Paragraph ID: {item['pid']}\n")
            if item['is_table']:
                f.write("**Source**: Table Cell\n")
            f.write(f"**Original**:\n{item['text']}\n\n")
            f.write("**Entities**:\n")
            
            if not item["matches"]:
                f.write("- [NONE]\n")
            else:
                for m in item["matches"]:
                    f.write(f"- [{m['type']}] {m['value']}\n")
            f.write("\n---\n\n")
            
    print(f"Draft annotations written to {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    
    random.seed(42)
    generate_sample(args.input, args.output)
